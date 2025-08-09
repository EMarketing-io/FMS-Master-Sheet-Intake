from typing import Dict, Any
from datetime import date
from docx import Document
import io


def generate_docx(
    summary_data: Dict[str, Any], company_name: str, meeting_date, mode: str = "full"
) -> io.BytesIO:
    """
    Create a .docx meeting document (full/MoM/Action) from structured summary data.
    (Flow and formatting unchanged.)
    """
    doc = Document()

    # Normalize date to dd-mm-YYYY if a date object is passed
    if isinstance(meeting_date, date):
        meeting_date = meeting_date.strftime("%d-%m-%Y")

    # Title
    if mode == "full":
        doc.add_heading(f"{company_name} Meeting Notes", level=0)
        doc.add_paragraph(f"Date: {meeting_date}", style="Heading 2").alignment = 2
    elif mode == "mom":
        doc.add_heading(f"{company_name} - MoM", level=0)
    elif mode == "action":
        doc.add_heading(f"{company_name} - Action Points", level=0)

    # 1) MoM
    if mode in ["full", "mom"]:
        doc.add_heading("1. Minutes of the Meeting (MoM)", level=1)
        for point in summary_data.get("mom", []):
            doc.add_paragraph(point.strip(), style="List Bullet")

    # 2) To-Do (only in full)
    if mode == "full":
        doc.add_heading("2. To-Do List", level=1)
        for task in summary_data.get("todo_list", []):
            doc.add_paragraph(task.strip(), style="List Bullet")

    # 3) Action Points
    if mode in ["full", "action"]:
        doc.add_heading("3. Action Points / Action Plan", level=1)
        section_titles = {
            "decision_made": "Key Decisions Made",
            "key_services_to_promote": "Key Services to Promote",
            "target_geography": "Target Geography",
            "budget_and_timeline": "Budget and Timeline",
            "lead_management_strategy": "Lead Management Strategy",
            "next_steps_and_ownership": "Next Steps and Ownership",
        }
        action_plan = summary_data.get("action_plan", {})
        for key, title in section_titles.items():
            items = action_plan.get(key, [])
            if items:
                doc.add_heading(title, level=2)
                for item in items:
                    doc.add_paragraph(item.strip(), style="List Bullet")

    # Export to BytesIO
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
