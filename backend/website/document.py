from docx import Document
from datetime import date
import io
import re


def generate_website_docx(summary_json, company_name: str, meeting_date) -> io.BytesIO:
    if isinstance(meeting_date, date):
        meeting_date = meeting_date.strftime("%d-%m-%Y")

    doc = Document()
    doc.add_heading(f"{company_name} Website Summary", level=0)
    p = doc.add_paragraph(f"Date: {meeting_date}")
    p.alignment = 2 

    for section in (summary_json.get("sections") or []):
        heading = (section.get("heading") or "").strip()
        content = (section.get("content") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)

        if content:
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("- "):
                    line = line[2:].strip()
                    para = doc.add_paragraph(style="List Bullet")
                    parts = re.split(r"(\*\*.*?\*\*)", line)
                    for part in parts:
                        run = para.add_run()
                        if part.startswith("**") and part.endswith("**"):
                            run.text = part[2:-2]
                            run.bold = True
                        else:
                            run.text = part
                else:
                    doc.add_paragraph(line)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out